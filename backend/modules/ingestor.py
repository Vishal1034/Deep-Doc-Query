import os
import json
import time
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional
from dotenv import load_dotenv

import chromadb
from sentence_transformers import SentenceTransformer

load_dotenv()

BASE_BACKEND_DIR = Path(__file__).resolve().parents[1]
DOCS_PATH = Path(os.getenv("DOCS_DIR", str(BASE_BACKEND_DIR / "docs")))
CHROMA_PATH = Path(os.getenv("CHROMA_PATH", str(BASE_BACKEND_DIR / "data" / "chroma")))
EMBEDDING_MODEL_NAME = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")

_embedding_model: Optional[SentenceTransformer] = None
_chroma_client: Optional[chromadb.PersistentClient] = None
_collection = None

def get_embedding_model() -> SentenceTransformer:
    global _embedding_model
    if _embedding_model is None:
        print(f"[Embeddings] Loading SentenceTransformer model: {EMBEDDING_MODEL_NAME}...")
        _embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
    return _embedding_model

def get_chroma_collection():
    global _chroma_client, _collection
    if _collection is None:
        CHROMA_PATH.mkdir(parents=True, exist_ok=True)
        _chroma_client = chromadb.PersistentClient(path=str(CHROMA_PATH))
        _collection = _chroma_client.get_or_create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"}
        )
    return _collection

def extract_text_from_file(file_path: Path) -> List[Dict[str, Any]]:
    """
    Extracts text and page/section metadata from supported file formats.
    Supported: .pdf, .docx, .txt, .md, .csv, .json, .py, .js, .ts, .html, .cpp, etc.
    """
    ext = file_path.suffix.lower()
    pages_data = []

    if ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(str(file_path))
            for page_idx, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages_data.append({
                        "text": text.strip(),
                        "page": page_idx + 1,
                        "section": f"Page {page_idx + 1}"
                    })
        except Exception as e:
            print(f"[Ingest Warning] Error reading PDF with pypdf: {e}, falling back to plain text read.")
            try:
                content = file_path.read_text(encoding="utf-8", errors="ignore")
                pages_data.append({"text": content, "page": 1, "section": "Full Document"})
            except Exception:
                pass

    elif ext == ".docx":
        try:
            import docx
            doc = docx.Document(str(file_path))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)
            text_content = "\n\n".join(full_text)
            if text_content:
                pages_data.append({"text": text_content, "page": 1, "section": "Document Content"})
        except Exception as e:
            print(f"[Ingest Warning] Error reading DOCX: {e}")

    elif ext == ".json":
        try:
            content = file_path.read_text(encoding="utf-8", errors="ignore")
            data = json.loads(content)
            formatted = json.dumps(data, indent=2)
            pages_data.append({"text": formatted, "page": 1, "section": "JSON Structure"})
        except Exception:
            content = file_path.read_text(encoding="utf-8", errors="ignore")
            pages_data.append({"text": content, "page": 1, "section": "JSON Content"})

    elif ext == ".csv":
        try:
            content = file_path.read_text(encoding="utf-8", errors="ignore")
            lines = content.splitlines()
            header = lines[0] if lines else ""
            batch_size = 20
            for i in range(1, len(lines), batch_size):
                batch = lines[i:i + batch_size]
                text = f"CSV Header: {header}\n" + "\n".join(batch)
                pages_data.append({
                    "text": text,
                    "page": (i // batch_size) + 1,
                    "section": f"Rows {i} to {min(i + batch_size, len(lines))}"
                })
            if not pages_data and content.strip():
                pages_data.append({"text": content, "page": 1, "section": "CSV Table"})
        except Exception:
            content = file_path.read_text(encoding="utf-8", errors="ignore")
            pages_data.append({"text": content, "page": 1, "section": "CSV Data"})

    else:
        # Markdown, Code, Plain text, etc.
        try:
            content = file_path.read_text(encoding="utf-8", errors="ignore")
            if content.strip():
                pages_data.append({"text": content.strip(), "page": 1, "section": "Content"})
        except Exception as e:
            print(f"[Ingest Warning] Error reading file {file_path.name}: {e}")

    return pages_data


def chunk_text(text: str, chunk_size: int = 600, chunk_overlap: int = 80) -> List[str]:
    """
    Intelligent character/token sliding chunker with paragraph and boundary preservation.
    """
    if not text:
        return []
    
    paragraphs = text.split("\n\n")
    chunks = []
    current_chunk = []
    current_len = 0

    for para in paragraphs:
        para_len = len(para)
        if current_len + para_len > chunk_size and current_chunk:
            combined = "\n\n".join(current_chunk).strip()
            if combined:
                chunks.append(combined)
            if chunk_overlap > 0 and len(current_chunk) > 1:
                current_chunk = [current_chunk[-1], para]
                current_len = len(current_chunk[-1]) + para_len
            else:
                current_chunk = [para]
                current_len = para_len
        else:
            current_chunk.append(para)
            current_len += para_len + 2

    if current_chunk:
        combined = "\n\n".join(current_chunk).strip()
        if combined:
            chunks.append(combined)

    refined_chunks = []
    for c in chunks:
        if len(c) > chunk_size * 1.5:
            start = 0
            while start < len(c):
                end = start + chunk_size
                refined_chunks.append(c[start:end].strip())
                start += chunk_size - chunk_overlap
        else:
            refined_chunks.append(c)

    return [c for c in refined_chunks if len(c.strip()) > 15]


def index_single_file(file_path: Path) -> Dict[str, Any]:
    """
    Indexes a single file into ChromaDB with deduplication and metadata.
    """
    collection = get_chroma_collection()
    model = get_embedding_model()

    filename = file_path.name
    file_size = file_path.stat().st_size if file_path.exists() else 0
    ext = file_path.suffix.lower()

    # 1. Remove existing chunks for this file if present (Deduplication)
    try:
        existing = collection.get(where={"source": filename})
        if existing and existing.get("ids"):
            collection.delete(ids=existing["ids"])
            print(f"[Ingest] Cleaned {len(existing['ids'])} previous chunks for {filename}")
    except Exception as e:
        print(f"[Ingest Note] Note on cleaning existing chunks: {e}")

    # 2. Extract text per section/page
    pages_data = extract_text_from_file(file_path)
    if not pages_data:
        return {"filename": filename, "chunks_indexed": 0, "status": "empty"}

    all_chunks = []
    all_metadatas = []
    all_ids = []

    global_chunk_idx = 0
    for page_item in pages_data:
        text = page_item["text"]
        page_num = page_item.get("page", 1)
        section = page_item.get("section", "General")

        text_chunks = chunk_text(text)
        for chunk in text_chunks:
            chunk_id = f"{filename}_chunk_{global_chunk_idx}_{int(time.time())}"
            metadata = {
                "source": filename,
                "file_type": ext,
                "file_size": file_size,
                "page": page_num,
                "section": section,
                "chunk_index": global_chunk_idx,
                "char_length": len(chunk),
                "timestamp": int(time.time())
            }
            all_chunks.append(chunk)
            all_metadatas.append(metadata)
            all_ids.append(chunk_id)
            global_chunk_idx += 1

    if not all_chunks:
        return {"filename": filename, "chunks_indexed": 0, "status": "no_chunks"}

    # 3. Generate embeddings
    embeddings = model.encode(all_chunks, show_progress_bar=False, normalize_embeddings=True).tolist()

    # 4. Insert into ChromaDB
    collection.add(
        ids=all_ids,
        documents=all_chunks,
        embeddings=embeddings,
        metadatas=all_metadatas
    )

    print(f"[Ingest] Indexed {filename}: {len(all_chunks)} chunks added to ChromaDB.")
    return {
        "filename": filename,
        "chunks_indexed": len(all_chunks),
        "file_size": file_size,
        "file_type": ext,
        "status": "indexed"
    }


def ingest_all_docs() -> Dict[str, Any]:
    """
    Scans the DOCS_PATH folder and indexes all files.
    """
    DOCS_PATH.mkdir(parents=True, exist_ok=True)
    supported_exts = {".pdf", ".docx", ".txt", ".md", ".csv", ".json", ".py", ".js", ".ts", ".html", ".cpp", ".c", ".java", ".go", ".rs", ".yml", ".yaml"}
    
    files = [f for f in DOCS_PATH.iterdir() if f.is_file() and f.suffix.lower() in supported_exts]
    results = []
    for file in files:
        res = index_single_file(file)
        results.append(res)
    
    return {
        "total_files": len(files),
        "details": results
    }


def list_indexed_documents() -> List[Dict[str, Any]]:
    """
    Returns a summarized list of all unique indexed documents in the database.
    """
    collection = get_chroma_collection()
    try:
        data = collection.get(include=["metadatas"])
        if not data or not data.get("metadatas"):
            return []

        doc_summary: Dict[str, Dict[str, Any]] = {}
        for meta in data["metadatas"]:
            if not meta or not meta.get("source"):
                continue
            source = meta["source"]
            if source not in doc_summary:
                doc_summary[source] = {
                    "filename": source,
                    "file_type": meta.get("file_type", ".txt"),
                    "file_size": meta.get("file_size", 0),
                    "chunk_count": 0,
                    "timestamp": meta.get("timestamp", int(time.time())),
                    "pages": set()
                }
            doc_summary[source]["chunk_count"] += 1
            if meta.get("page"):
                doc_summary[source]["pages"].add(meta["page"])

        result = []
        for src, item in doc_summary.items():
            result.append({
                "filename": item["filename"],
                "file_type": item["file_type"],
                "file_size": item["file_size"],
                "chunk_count": item["chunk_count"],
                "total_pages": len(item["pages"]) if item["pages"] else 1,
                "timestamp": item["timestamp"]
            })
        return sorted(result, key=lambda x: x["filename"])
    except Exception as e:
        print(f"[Ingest Error] Error listing documents: {e}")
        return []


def get_document_chunks(filename: str) -> List[Dict[str, Any]]:
    """
    Returns all chunks and metadata for a specific indexed document.
    """
    collection = get_chroma_collection()
    try:
        data = collection.get(where={"source": filename}, include=["documents", "metadatas"])
        if not data or not data.get("documents"):
            return []
        
        chunks = []
        for doc_id, text, meta in zip(data["ids"], data["documents"], data["metadatas"]):
            chunks.append({
                "id": doc_id,
                "text": text,
                "metadata": meta
            })
        return sorted(chunks, key=lambda x: x["metadata"].get("chunk_index", 0))
    except Exception as e:
        print(f"[Ingest Error] Error fetching chunks for {filename}: {e}")
        return []


def delete_document_from_db(filename: str) -> bool:
    """
    Deletes all chunks of a document from ChromaDB and removes the physical file if present.
    """
    collection = get_chroma_collection()
    deleted = False
    try:
        existing = collection.get(where={"source": filename})
        if existing and existing.get("ids"):
            collection.delete(ids=existing["ids"])
            deleted = True
            print(f"[Ingest] Deleted {len(existing['ids'])} chunks for {filename}")
    except Exception as e:
        print(f"[Ingest Error] Error deleting from ChromaDB: {e}")

    # Remove physical file from docs folder if exists
    try:
        physical_path = DOCS_PATH / filename
        if physical_path.exists():
            physical_path.unlink()
            deleted = True
    except Exception as e:
        print(f"[Ingest Error] Error deleting physical file: {e}")

    return deleted


def get_collection_statistics() -> Dict[str, Any]:
    """
    Returns total document counts, chunk counts, and embedding model information.
    """
    collection = get_chroma_collection()
    try:
        count = collection.count()
        docs = list_indexed_documents()
        total_size = sum(d.get("file_size", 0) for d in docs)
        return {
            "total_documents": len(docs),
            "total_chunks": count,
            "total_size_bytes": total_size,
            "embedding_model": EMBEDDING_MODEL_NAME,
            "chroma_path": str(CHROMA_PATH)
        }
    except Exception as e:
        return {
            "total_documents": 0,
            "total_chunks": 0,
            "total_size_bytes": 0,
            "embedding_model": EMBEDDING_MODEL_NAME,
            "error": str(e)
        }
